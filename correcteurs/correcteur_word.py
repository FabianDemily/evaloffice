"""Correcteur Word générique, piloté par config JSON.

Chaque critère est identifié par son type et ses ancres textuelles.
La correction est tolérante aux variantes linguistiques (FR/EN) via styles_multilingues.
"""

import unicodedata
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .styles_multilingues import canonique, est_style

STATUT_REUSSI = "reussi"
STATUT_ECHOUE = "echoue"

FEUILLE_GARDE = None  # Word n'a pas d'onglets


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normaliser_texte(s):
    """Minuscules sans accents pour comparaisons tolérantes."""
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return s.lower()


def _trouver_paragraphe(doc, ancre):
    """Retourne le paragraphe correspondant à l'ancre (insensible à la casse/accents).

    Privilégie la correspondance exacte pour éviter de confondre l'ancre réelle
    avec les paragraphes de consignes qui la citent entre guillemets.
    """
    ancre_n = _normaliser_texte(ancre)
    # 1. Correspondance exacte
    for p in doc.paragraphs:
        if _normaliser_texte(p.text) == ancre_n:
            return p
    # 2. Repli : contient l'ancre
    for p in doc.paragraphs:
        if ancre_n in _normaliser_texte(p.text):
            return p
    return None


def _texte_document(doc):
    """Retourne tout le texte du document en minuscules sans accents."""
    return _normaliser_texte(" ".join(p.text for p in doc.paragraphs))


def _texte_entete(doc):
    try:
        return _normaliser_texte(
            " ".join(p.text for p in doc.sections[0].header.paragraphs)
        )
    except Exception:
        return ""


def _texte_pied(doc):
    try:
        return _normaliser_texte(
            " ".join(p.text for p in doc.sections[0].footer.paragraphs)
        )
    except Exception:
        return ""


def _compter_notes_de_bas_de_page(doc):
    """Compte les notes de bas de page utilisateur (hors séparateurs système)."""
    FOOTNOTES_REL = (
        "http://schemas.openxmlformats.org/officeDocument/2006/"
        "relationships/footnotes"
    )
    try:
        from lxml import etree
        for rel in doc.part.rels.values():
            if rel.reltype == FOOTNOTES_REL:
                root = etree.fromstring(rel.target_part.blob)
                notes = root.findall(qn("w:footnote"))
                # Les id -1 et 0 sont les séparateurs système
                user_notes = [
                    n for n in notes
                    if n.get(qn("w:id")) not in ("-1", "0")
                    and n.get(qn("w:type")) is None
                ]
                return len(user_notes)
    except Exception:
        pass
    return 0


def _a_saut_de_page(doc, ancre):
    """Vérifie si un saut de page (manuel ou par propriété) précède le paragraphe ancré."""
    # Utilise _trouver_paragraphe pour éviter de matcher les paragraphes de consignes
    # qui citent l'ancre entre guillemets plutôt que l'ancre réelle dans le corps.
    para = _trouver_paragraphe(doc, ancre)
    if para is None:
        return False

    # Retrouver l'index en comparant les éléments XML (doc.paragraphs retourne
    # de nouveaux objets Python à chaque appel, donc .index() échoue).
    paragraphes = doc.paragraphs
    idx = next((i for i, p in enumerate(paragraphes) if p._p is para._p), None)
    if idx is None:
        return False

    # Vérifie un saut dans le paragraphe lui-même (w:br type="page")
    for br in para._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True

    # Vérifie la propriété pageBreakBefore
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        pbr = pPr.find(qn("w:pageBreakBefore"))
        if pbr is not None and pbr.get(qn("w:val"), "1") != "0":
            return True

    # Vérifie si les 1 ou 2 paragraphes précédents contiennent un saut de page
    # (Word insère souvent le saut dans un paragraphe vide juste avant)
    for offset in (1, 2):
        if idx >= offset:
            prev = paragraphes[idx - offset]
            for br in prev._p.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    return True

    return False


def _a_champ_toc(doc):
    """Vérifie si le document contient un champ TOC ou des paragraphes de style TDM."""
    # Recherche dans le XML du corps
    corps_xml = doc.element.body.xml
    if "TOC" in corps_xml or "\\o" in corps_xml:
        return True
    # Repli : cherche des paragraphes de style TDM
    for p in doc.paragraphs:
        if canonique(p.style.style_id) in ("toc1", "toc2", "toc3"):
            return True
        if canonique(p.style.name) in ("toc1", "toc2", "toc3"):
            return True
    return False


# ---------------------------------------------------------------------------
# Fonctions de correction par type
# ---------------------------------------------------------------------------

def _corriger_mise_en_page(doc, critere):
    points_max = critere.get("points", 1)
    orientation_attendue = critere.get("orientation", "paysage")
    marge_attendue_cm = critere.get("marge_cm", 2.0)
    tolerance_cm = 0.3  # tolérance de ±3 mm

    section = doc.sections[0]
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    details = []
    pts = 0
    pts_par = points_max / 2

    # Orientation
    est_paysage = section.page_width > section.page_height
    if orientation_attendue == "paysage" and est_paysage:
        pts += pts_par
        details.append(f"Orientation paysage : +{pts_par:.2f} pt")
    elif orientation_attendue == "portrait" and not est_paysage:
        pts += pts_par
        details.append(f"Orientation portrait : +{pts_par:.2f} pt")
    else:
        orientation_trouvee = "paysage" if est_paysage else "portrait"
        details.append(
            f"Orientation incorrecte : {orientation_trouvee} trouvée, "
            f"{orientation_attendue} attendue (0 pt)"
        )

    # Marges (on vérifie la marge gauche comme indicateur)
    from docx.shared import Cm as _Cm
    marge_gauche_cm = section.left_margin.cm if section.left_margin else 999
    if abs(marge_gauche_cm - marge_attendue_cm) <= tolerance_cm:
        pts += pts_par
        details.append(f"Marges {marge_gauche_cm:.1f} cm : +{pts_par:.2f} pt")
    else:
        details.append(
            f"Marges incorrectes : {marge_gauche_cm:.1f} cm trouvés, "
            f"{marge_attendue_cm} cm attendus (0 pt)"
        )

    return points_max, pts, details


def _corriger_structure_hierarchique(doc, critere):
    points_max = critere.get("points", 1)
    paragraphes = critere.get("paragraphes", [])
    if not paragraphes:
        return points_max, 0, ["Aucun paragraphe à vérifier"]

    pts_par = points_max / len(paragraphes)
    pts = 0.0
    details = []

    for spec in paragraphes:
        ancre = spec["ancre"]
        style_attendu = spec["style"]  # clé canonique ex: "heading1"
        para = _trouver_paragraphe(doc, ancre)
        if para is None:
            details.append(f"Paragraphe « {ancre[:30]} » : introuvable (0 pt)")
            continue
        if est_style(para, style_attendu):
            pts += pts_par
            details.append(f"« {ancre[:30]} » : style correct (+{pts_par:.2f} pt)")
        else:
            style_trouve = para.style.name
            details.append(
                f"« {ancre[:30]} » : style incorrect ({style_trouve} trouvé, "
                f"{style_attendu} attendu) (0 pt)"
            )

    return points_max, pts, details


def _corriger_styles_paragraphe(doc, critere):
    points_max = critere.get("points", 1)
    style_attendu = critere.get("style_attendu", "body_text")
    ancres = critere.get("ancres", [])
    taille_attendue_pt = critere.get("taille_pt")        # ex: 13
    justification_attendue = critere.get("justifie", False)
    tolerance_pt = 1

    if not ancres:
        return points_max, 0, ["Aucun paragraphe à vérifier"]

    # Sous-critères : application du style + (optionnel) taille + justification
    nb_sous = 1 + (1 if taille_attendue_pt else 0) + (1 if justification_attendue else 0)
    pts_par_sous = points_max / nb_sous
    pts = 0.0
    details = []

    # --- 1. Style appliqué aux paragraphes ---
    nb_ok = sum(
        1 for ancre in ancres
        if (p := _trouver_paragraphe(doc, ancre)) is not None and est_style(p, style_attendu)
    )
    if nb_ok == len(ancres):
        pts += pts_par_sous
        details.append(f"Style « Corps de texte » appliqué ({nb_ok}/{len(ancres)} paragraphes) : +{pts_par_sous:.2f} pt")
    elif nb_ok > 0:
        partiel = pts_par_sous * nb_ok / len(ancres)
        pts += partiel
        details.append(f"Style partiel ({nb_ok}/{len(ancres)} paragraphes) : +{partiel:.2f} pt")
    else:
        for ancre in ancres:
            p = _trouver_paragraphe(doc, ancre)
            nom = p.style.name if p else "introuvable"
            details.append(f"Style incorrect sur « {ancre[:30]}... » : {nom} trouvé (0 pt)")

    # --- 2. Taille de police modifiée dans la définition du style ---
    if taille_attendue_pt:
        try:
            style_obj = None
            for s in doc.styles:
                if canonique(s.style_id) == style_attendu or canonique(s.name) == style_attendu:
                    style_obj = s
                    break
            taille_trouvee = style_obj.font.size.pt if (style_obj and style_obj.font.size) else None
            if taille_trouvee and abs(taille_trouvee - taille_attendue_pt) <= tolerance_pt:
                pts += pts_par_sous
                details.append(f"Taille du style : {taille_trouvee:.0f} pt : +{pts_par_sous:.2f} pt")
            else:
                val = f"{taille_trouvee:.0f} pt" if taille_trouvee else "non définie"
                details.append(f"Taille du style incorrecte : {val} (attendu {taille_attendue_pt} pt) (0 pt)")
        except Exception:
            details.append("Taille du style : impossible à lire (0 pt)")

    # --- 3. Justification du texte dans la définition du style ---
    if justification_attendue:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        try:
            style_obj = None
            for s in doc.styles:
                if canonique(s.style_id) == style_attendu or canonique(s.name) == style_attendu:
                    style_obj = s
                    break
            align = style_obj.paragraph_format.alignment if style_obj else None
            if align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                pts += pts_par_sous
                details.append(f"Justification du style : correcte : +{pts_par_sous:.2f} pt")
            else:
                nom_align = align.name if align is not None else "non définie"
                details.append(f"Justification du style : {nom_align} (justifié attendu) (0 pt)")
        except Exception:
            details.append("Justification du style : impossible à lire (0 pt)")

    return points_max, pts, details


def _corriger_mise_en_forme_caracteres(doc, critere):
    points_max = critere.get("points", 1)
    ancre = critere.get("ancre_paragraphe", "")
    mot_cible = critere.get("mot_cible", "")
    gras_attendu = critere.get("gras", False)
    taille_attendue = critere.get("taille_pt")
    couleur_attendue = critere.get("couleur_hex", "").upper().replace("#", "")
    tolerance_pt = 1

    para = _trouver_paragraphe(doc, ancre)
    if para is None:
        return points_max, 0, [f"Paragraphe ancre introuvable (0 pt)"]

    # Cherche le run contenant le mot cible
    mot_n = _normaliser_texte(mot_cible)
    run_cible = None
    for run in para.runs:
        if mot_n in _normaliser_texte(run.text):
            run_cible = run
            break

    if run_cible is None:
        return points_max, 0, [f"Mot « {mot_cible} » introuvable dans le paragraphe (0 pt)"]

    nb_criteres = sum([1, gras_attendu, bool(taille_attendue), bool(couleur_attendue)])
    pts_par = points_max / nb_criteres
    pts = pts_par  # point pour avoir trouvé le mot
    details = [f"Mot « {mot_cible} » trouvé : +{pts_par:.2f} pt"]

    if gras_attendu:
        if run_cible.bold:
            pts += pts_par
            details.append(f"Gras présent : +{pts_par:.2f} pt")
        else:
            details.append("Gras manquant (0 pt)")

    if taille_attendue:
        taille_trouvee = run_cible.font.size.pt if run_cible.font.size else None
        if taille_trouvee and abs(taille_trouvee - taille_attendue) <= tolerance_pt:
            pts += pts_par
            details.append(f"Taille {taille_trouvee:.0f} pt : +{pts_par:.2f} pt")
        else:
            details.append(
                f"Taille incorrecte : {taille_trouvee} pt trouvé, "
                f"{taille_attendue} pt attendu (0 pt)"
            )

    if couleur_attendue:
        try:
            couleur_trouvee = str(run_cible.font.color.rgb).upper()
        except Exception:
            couleur_trouvee = ""
        if couleur_trouvee == couleur_attendue:
            pts += pts_par
            details.append(f"Couleur #{couleur_trouvee} : +{pts_par:.2f} pt")
        else:
            details.append(
                f"Couleur incorrecte : #{couleur_trouvee} trouvée, "
                f"#{couleur_attendue} attendue (0 pt)"
            )

    return points_max, pts, details


def _corriger_rechercher_remplacer(doc, critere):
    points_max = critere.get("points", 1)
    terme_ancien = _normaliser_texte(critere.get("terme_ancien", ""))
    terme_nouveau = _normaliser_texte(critere.get("terme_nouveau", ""))
    texte = _texte_document(doc)
    details = []
    pts = 0

    ancien_present = terme_ancien in texte
    nouveau_present = terme_nouveau in texte

    if not ancien_present and nouveau_present:
        pts = points_max
        details.append(
            f"Remplacement effectué : « {critere['terme_nouveau']} » présent, "
            f"« {critere['terme_ancien']} » absent (+{points_max:.2f} pt)"
        )
    elif ancien_present and nouveau_present:
        pts = points_max / 2
        details.append(
            f"Remplacement partiel : « {critere['terme_ancien']} » encore présent "
            f"dans le document (+{pts:.2f} pt)"
        )
    elif not nouveau_present:
        details.append(
            f"Remplacement non effectué : "
            f"« {critere['terme_nouveau']} » absent (0 pt)"
        )

    return points_max, pts, details


def _corriger_saisie_texte(doc, critere):
    """Vérifie la saisie de caractères spéciaux depuis une image (non copiable).

    Deux sous-critères (points_max / 2 chacun) :
      1. Tous les caractères spéciaux sont présents dans la zone de saisie
      2. La séquence complète est exacte
    """
    points_max = critere.get("points", 1)
    ancre_zone = critere.get("ancre_zone", "Zone de saisie")
    sequence_attendue = critere.get("sequence_attendue", "")
    chars_speciaux = critere.get("chars_speciaux", [])

    # Repli vers l'ancien mode mots-clés si la config est ancienne
    if not sequence_attendue:
        mots_cles = [_normaliser_texte(m) for m in critere.get("mots_cles", [])]
        texte = _texte_document(doc)
        pts = sum(1 for m in mots_cles if m in texte) / max(len(mots_cles), 1) * points_max
        details = [f"Mot-clé « {m} » {'trouvé' if m in texte else 'absent'}" for m in mots_cles]
        return points_max, pts, details

    pts_par = points_max / 2
    pts = 0.0
    details = []

    # Trouver le texte saisi dans le tableau de saisie qui suit l'ancre
    para_zone = _trouver_paragraphe(doc, ancre_zone)
    texte_saisi = ""
    if para_zone is not None:
        # Chercher le premier tableau dont l'élément XML suit l'ancre dans le corps
        body = doc.element.body
        children = list(body)
        try:
            ancre_idx = next(i for i, el in enumerate(children) if el is para_zone._p)
        except StopIteration:
            ancre_idx = -1

        from docx.oxml.ns import qn as _qn
        for el in children[ancre_idx + 1:]:
            if el.tag == _qn("w:tbl"):
                # Tableau trouvé : lire cell(1, 0) = zone de frappe étudiant
                from docx.table import Table
                tbl = Table(el, doc)
                try:
                    texte_saisi = tbl.cell(1, 0).text.strip()
                except Exception:
                    pass
                break
            # S'arrêter si on rencontre un autre paragraphe non vide (hors ancre)
            if el.tag == _qn("w:p"):
                txt = "".join(r.text or "" for r in el.findall(f".//{_qn('w:t')}")).strip()
                if txt and txt != ancre_zone:
                    break

    if not texte_saisi:
        details.append("Zone de saisie vide ou introuvable (0 pt)")
        details.append("Séquence exacte : non évaluable (0 pt)")
        return points_max, 0, details

    # --- Sous-critère 1 : caractères spéciaux tous présents ---
    chars_presents = [c for c in chars_speciaux if c in texte_saisi]
    chars_absents  = [c for c in chars_speciaux if c not in texte_saisi]
    if not chars_absents:
        pts += pts_par
        details.append(f"Caractères spéciaux {chars_speciaux} tous présents : +{pts_par:.2f} pt")
    elif chars_presents:
        partiel = pts_par * len(chars_presents) / len(chars_speciaux)
        pts += partiel
        details.append(
            f"Caractères spéciaux partiels — présents : {chars_presents}, "
            f"absents : {chars_absents} : +{partiel:.2f} pt"
        )
    else:
        details.append(f"Aucun caractère spécial trouvé {chars_speciaux} (0 pt)")

    # --- Sous-critère 2 : séquence exacte (tolérance sur les espaces internes) ---
    # Un espace visuel entre deux caractères spéciaux (ex: "}  [") peut être dû
    # au rendu de la police dans l'image — on compare sans espaces internes.
    saisi_sans_espaces = texte_saisi.replace(" ", "")
    attendu_sans_espaces = sequence_attendue.replace(" ", "")
    if saisi_sans_espaces == attendu_sans_espaces:
        pts += pts_par
        details.append(f"Séquence correcte {sequence_attendue!r} : +{pts_par:.2f} pt")
    else:
        details.append(
            f"Séquence incorrecte : {texte_saisi!r} saisi, "
            f"{sequence_attendue!r} attendu (0 pt)"
        )

    return points_max, pts, details


def _corriger_stabilite_mise_en_page(doc, critere):
    points_max = critere.get("points", 1)
    ancre = critere.get("ancre_saut", "")
    ok = _a_saut_de_page(doc, ancre)
    if ok:
        details = [f"Saut de page détecté avant « {ancre[:40]} » : +{points_max:.2f} pt"]
    else:
        details = [f"Saut de page absent avant « {ancre[:40]} » (0 pt)"]
    return points_max, (points_max if ok else 0), details


def _corriger_image(doc, critere):
    points_max = critere.get("points", 1)
    largeur_min = critere.get("largeur_min_cm", 0)
    largeur_max = critere.get("largeur_max_cm", 999)

    shapes = doc.inline_shapes
    if not shapes:
        return points_max, 0, ["Aucune image détectée (0 pt)"]

    from docx.shared import Cm
    details = [f"{len(shapes)} image(s) détectée(s)"]
    pts_par = points_max / 2
    pts = pts_par  # présence

    img = shapes[0]
    largeur_cm = img.width.cm if img.width else 0
    if largeur_min <= largeur_cm <= largeur_max:
        pts += pts_par
        details.append(f"Largeur {largeur_cm:.1f} cm (attendu {largeur_min}–{largeur_max} cm) : +{pts_par:.2f} pt")
    else:
        details.append(
            f"Largeur {largeur_cm:.1f} cm hors plage {largeur_min}–{largeur_max} cm (0 pt)"
        )

    return points_max, pts, details


def _corriger_table_des_matieres(doc, critere):
    points_max = critere.get("points", 1)
    ok = _a_champ_toc(doc)
    if ok:
        details = [f"Table des matières détectée : +{points_max:.2f} pt"]
    else:
        details = ["Table des matières absente (0 pt)"]
    return points_max, (points_max if ok else 0), details


def _corriger_entete_pied_page(doc, critere):
    points_max = critere.get("points", 1)
    mots_entete = [_normaliser_texte(m) for m in critere.get("mots_cles_entete", [])]
    mots_pied = [_normaliser_texte(m) for m in critere.get("mots_cles_pied", [])]

    pts_par = points_max / 2
    pts = 0.0
    details = []

    texte_entete = _texte_entete(doc)
    entete_ok = all(m in texte_entete for m in mots_entete) if mots_entete else bool(texte_entete.strip())
    if entete_ok:
        pts += pts_par
        details.append(f"En-tête renseigné : +{pts_par:.2f} pt")
    else:
        manquants = [m for m in mots_entete if m not in texte_entete]
        details.append(f"En-tête incomplet — éléments manquants : {', '.join(manquants)} (0 pt)")

    texte_pied = _texte_pied(doc)
    pied_ok = all(m in texte_pied for m in mots_pied) if mots_pied else bool(texte_pied.strip())
    if pied_ok:
        pts += pts_par
        details.append(f"Pied de page renseigné : +{pts_par:.2f} pt")
    else:
        manquants = [m for m in mots_pied if m not in texte_pied]
        details.append(f"Pied de page incomplet — éléments manquants : {', '.join(manquants)} (0 pt)")

    return points_max, pts, details


def _corriger_note_de_bas_de_page(doc, critere):
    points_max = critere.get("points", 1)
    nb_min = critere.get("nb_notes_min", 1)
    nb = _compter_notes_de_bas_de_page(doc)
    if nb >= nb_min:
        details = [f"{nb} note(s) de bas de page détectée(s) : +{points_max:.2f} pt"]
        return points_max, points_max, details
    else:
        details = [f"Aucune note de bas de page détectée (0 pt)"]
        return points_max, 0, details


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_DISPATCH = {
    "mise_en_page":              lambda doc, critere: _corriger_mise_en_page(doc, critere),
    "structure_hierarchique":    lambda doc, critere: _corriger_structure_hierarchique(doc, critere),
    "styles_paragraphe":         lambda doc, critere: _corriger_styles_paragraphe(doc, critere),
    "mise_en_forme_caracteres":  lambda doc, critere: _corriger_mise_en_forme_caracteres(doc, critere),
    "rechercher_remplacer":      lambda doc, critere: _corriger_rechercher_remplacer(doc, critere),
    "saisie_texte":              lambda doc, critere: _corriger_saisie_texte(doc, critere),
    "stabilite_mise_en_page":    lambda doc, critere: _corriger_stabilite_mise_en_page(doc, critere),
    "image":                     lambda doc, critere: _corriger_image(doc, critere),
    "table_des_matieres":        lambda doc, critere: _corriger_table_des_matieres(doc, critere),
    "entete_pied_page":          lambda doc, critere: _corriger_entete_pied_page(doc, critere),
    "note_de_bas_de_page":       lambda doc, critere: _corriger_note_de_bas_de_page(doc, critere),
}


def corriger_critere_word(doc, critere):
    type_critere = critere["type"]
    if type_critere not in _DISPATCH:
        raise ValueError(f"Type de critère Word inconnu : {type_critere}")
    return _DISPATCH[type_critere](doc, critere)


def _extraire_metadonnees_docx(doc):
    """Extrait les métadonnées de traçabilité d'un document Word."""
    props = doc.core_properties
    creator = str(props.author or "").strip()
    last_modified_by = str(props.last_modified_by or "").strip()
    created = props.created
    modified = props.modified

    duree_minutes = None
    if created and modified and modified >= created:
        duree_minutes = round((modified - created).total_seconds() / 60, 1)

    alertes = []
    if duree_minutes is not None and duree_minutes < 5:
        alertes.append(f"Durée de travail très courte : {duree_minutes} min")

    return {
        "creator": creator or "—",
        "last_modified_by": last_modified_by or "—",
        "created": created.strftime("%Y-%m-%d %H:%M") if created else "—",
        "modified": modified.strftime("%Y-%m-%d %H:%M") if modified else "—",
        "duree_minutes": duree_minutes,
        "alertes": alertes,
    }


def corriger_copie_word(chemin_fichier, config):
    """Corrige une copie Word selon la config et retourne le résultat structuré."""
    doc = Document(str(chemin_fichier))

    # Identité depuis le nom de fichier
    parts = Path(chemin_fichier).stem.split("_")
    if len(parts) >= 2:
        identite = {"nom": parts[0], "prenom": parts[1]}
    else:
        identite = {"nom": Path(chemin_fichier).stem, "prenom": ""}

    resultat = {
        "identite": identite,
        "session": config.get("session"),
        "annee": config.get("annee"),
        "module": config.get("module"),
        "variante": config.get("variante"),
        "watermark": None,
        "metadonnees": _extraire_metadonnees_docx(doc),
        "exercices": [],
        "points_obtenus": 0.0,
        "points_max": 0.0,
    }

    for exercice in config["exercices"]:
        ex_res = {
            "id": exercice["id"],
            "titre": exercice["titre"],
            "criteres": [],
            "points_obtenus": 0.0,
            "points_max": 0.0,
        }
        for critere in exercice["criteres"]:
            pts_max, pts_obt, details = corriger_critere_word(doc, critere)
            statut = STATUT_REUSSI if pts_obt >= pts_max - 1e-6 else STATUT_ECHOUE
            ex_res["criteres"].append({
                "id": critere["id"],
                "competence": critere.get("competence", critere.get("description", "")),
                "description": critere.get("description", ""),
                "points_max": pts_max,
                "points_obtenus": pts_obt,
                "statut": statut,
                "details": details,
            })
            ex_res["points_obtenus"] += pts_obt
            ex_res["points_max"] += pts_max

        resultat["exercices"].append(ex_res)
        resultat["points_obtenus"] += ex_res["points_obtenus"]
        resultat["points_max"] += ex_res["points_max"]

    resultat["bareme_total"] = config.get("bareme_total", resultat["points_max"])
    return resultat
