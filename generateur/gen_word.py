"""Génération d'épreuves Word (fichier étudiant + corrigé) et de leur config JSON.

Le document généré est un rapport thématique unique. Chaque compétence active
ajoute un élément structurel au document ; la correction utilise les ancres
(textes de repère) stockées dans la config pour localiser chaque élément.
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RACINE = Path(__file__).parent.parent
THEMES_PATH = RACINE / "data" / "themes.json"
COMPETENCES_WORD_PATH = RACINE / "data" / "competences_word.json"

# Contenu thématique par contexte
CONTEXTES = {
    "librairie": {
        "titre":          "Rapport sur la librairie",
        "terme_remplacer": "ouvrage",
        "terme_nouveau":   "livre",
        "mot_formater":    "essentielle",
        "terme_note":      "catalogue",
        "intro_body": (
            "La librairie joue un rôle essentielle dans la diffusion de la culture. "
            "Chaque ouvrage proposé répond à une demande précise des lecteurs."
        ),
        "donnees_body": (
            "L'analyse des ventes révèle que certains ouvrages dominent le marché. "
            "La gestion du catalogue permet d'optimiser les commandes."
        ),
        "saisie_texte":   "La librairie est un lieu d'échange et de découverte culturelle.",
        "mots_cles_saisie": ["librairie", "échange", "culturelle"],
    },
    "pharmacie": {
        "titre":          "Rapport sur la pharmacie",
        "terme_remplacer": "médicament",
        "terme_nouveau":   "produit pharmaceutique",
        "mot_formater":    "primordiale",
        "terme_note":      "ordonnance",
        "intro_body": (
            "La pharmacie remplit une fonction primordiale dans le système de santé. "
            "Chaque médicament délivré est soumis à un contrôle rigoureux."
        ),
        "donnees_body": (
            "L'analyse des prescriptions montre une évolution des besoins. "
            "La gestion de l'ordonnance garantit la sécurité du patient."
        ),
        "saisie_texte":   "La pharmacie assure la sécurité et l'efficacité des soins.",
        "mots_cles_saisie": ["pharmacie", "sécurité", "soins"],
    },
    "sport": {
        "titre":          "Rapport sur les activités sportives",
        "terme_remplacer": "compétition",
        "terme_nouveau":   "épreuve sportive",
        "mot_formater":    "fondamentale",
        "terme_note":      "performance",
        "intro_body": (
            "Le sport occupe une place fondamentale dans notre société. "
            "Chaque compétition représente un défi pour les athlètes participants."
        ),
        "donnees_body": (
            "Les résultats obtenus lors des dernières compétitions sont analysés ici. "
            "La notion de performance est au cœur de chaque discipline."
        ),
        "saisie_texte":   "Le sport favorise le dépassement de soi et l'esprit d'équipe.",
        "mots_cles_saisie": ["sport", "dépassement", "équipe"],
    },
    "kinesitherapie": {
        "titre":          "Rapport sur la kinésithérapie",
        "terme_remplacer": "patient",
        "terme_nouveau":   "personne soignée",
        "mot_formater":    "thérapeutique",
        "terme_note":      "rééducation",
        "intro_body": (
            "La kinésithérapie propose une approche thérapeutique personnalisée. "
            "Chaque patient bénéficie d'un programme de soins adapté à ses besoins."
        ),
        "donnees_body": (
            "Le suivi des patients montre des progrès significatifs. "
            "Le processus de rééducation mobilise des techniques variées."
        ),
        "saisie_texte":   "La kinésithérapie améliore la qualité de vie des patients.",
        "mots_cles_saisie": ["kinésithérapie", "qualité", "patients"],
    },
    "centre_readaptation": {
        "titre":          "Rapport sur le centre de réadaptation",
        "terme_remplacer": "résident",
        "terme_nouveau":   "personne accueillie",
        "mot_formater":    "indispensable",
        "terme_note":      "protocole",
        "intro_body": (
            "Le centre de réadaptation offre un accompagnement indispensable. "
            "Chaque résident dispose d'un plan de soins individualisé."
        ),
        "donnees_body": (
            "L'évaluation des résidents révèle des besoins spécifiques. "
            "L'application du protocole garantit la qualité de la prise en charge."
        ),
        "saisie_texte":   "Le centre accompagne les résidents vers une vie autonome.",
        "mots_cles_saisie": ["centre", "autonome", "résidents"],
    },
}


def charger_competences_word():
    return json.loads(COMPETENCES_WORD_PATH.read_text(encoding="utf-8"))


def charger_themes():
    return json.loads(THEMES_PATH.read_text(encoding="utf-8"))


def _points(competences_par_id, comp_id):
    return competences_par_id[comp_id].get("points_defaut", 1)


# ---------------------------------------------------------------------------
# Helpers XML
# ---------------------------------------------------------------------------

def _inserer_saut_de_page(para):
    """Insère un saut de page manuel au début d'un paragraphe."""
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    para._p.insert(0, r)


def _inserer_champ_toc(para):
    """Insère un champ TOC dans un paragraphe."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r1 = OxmlElement("w:r")
    r1.append(fld_begin)
    r2 = OxmlElement("w:r")
    r2.append(instr)
    r3 = OxmlElement("w:r")
    r3.append(fld_sep)
    r4 = OxmlElement("w:r")
    r4.append(fld_end)

    for r in (r1, r2, r3, r4):
        para._p.append(r)


def _set_run_format(run, gras=False, taille_pt=None, couleur_hex=None, italique=False):
    if gras:
        run.bold = True
    if italique:
        run.italic = True
    if taille_pt:
        run.font.size = Pt(taille_pt)
    if couleur_hex:
        r, g, b = int(couleur_hex[0:2], 16), int(couleur_hex[2:4], 16), int(couleur_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _ajouter_tableau_saisie(doc, texte_prefill=""):
    """Insère un tableau 2 lignes pour la zone de saisie étudiant.

    Ligne 0 : label grisé (instruction).
    Ligne 1 : cellule de frappe, fond jaune pâle, grande hauteur, bordure bleue.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    from lxml import etree

    def _set_cell_bg(cell, couleur_hex):
        """Fond de cellule (couleur hex sans #)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:color"), "auto")
        shd.set(_qn("w:fill"), couleur_hex)
        tcPr.append(shd)

    def _set_table_borders(table, couleur_hex, taille_pt=12):
        """Bordure extérieure du tableau."""
        tbl = table._tbl
        tblPr = tbl.find(_qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        sz = str(taille_pt * 8)  # unité = 1/8 pt
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(_qn("w:val"), "single")
            el.set(_qn("w:sz"), sz)
            el.set(_qn("w:space"), "0")
            el.set(_qn("w:color"), couleur_hex)
            tblBorders.append(el)
        tblPr.append(tblBorders)

    def _set_cell_height(row, hauteur_twips):
        """Hauteur minimale d'une ligne (1 cm ≈ 567 twips)."""
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(_qn("w:val"), str(hauteur_twips))
        trHeight.set(_qn("w:hRule"), "atLeast")
        trPr.append(trHeight)

    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"

    # Ligne 0 — label instructionnel
    cell_label = table.cell(0, 0)
    _set_cell_bg(cell_label, "D9D9D9")  # gris clair
    p_label = cell_label.paragraphs[0]
    run_label = p_label.add_run("▶  Tapez votre réponse ici :")
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_label.font.italic = True

    # Ligne 1 — zone de frappe
    cell_saisie = table.cell(1, 0)
    _set_cell_bg(cell_saisie, "FFFCDC")  # jaune pâle
    _set_cell_height(table.rows[1], 1134)  # ~2 cm
    p_saisie = cell_saisie.paragraphs[0]
    if texte_prefill:
        run_s = p_saisie.add_run(texte_prefill)
        run_s.font.size = Pt(14)
    else:
        p_saisie.paragraph_format.space_before = Pt(6)

    _set_table_borders(table, "3A6BC4", taille_pt=2)  # bleu, 2 pt


def _generer_image_sequence(sequence):
    """Génère un PNG (BytesIO) affichant la séquence en caractères lisibles.

    L'image est non sélectionnable dans Word — l'étudiant doit taper au clavier.
    """
    import io
    from PIL import Image, ImageDraw, ImageFont

    FOND = (255, 252, 220)       # jaune pâle
    BORDURE = (180, 140, 0)
    TEXTE = (30, 30, 30)

    TAILLE_PT = 36
    PADDING = 20

    POLICES_CANDIDATS = [
        "/System/Library/Fonts/Courier New Bold.ttf",          # macOS
        "/System/Library/Fonts/Supplemental/Courier New.ttf",
        "/Library/Fonts/Courier New Bold.ttf",
        "C:/Windows/Fonts/courbd.ttf",                          # Windows
        "C:/Windows/Fonts/cour.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationMono-Bold.ttf",  # Linux
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf",
    ]
    font = None
    for chemin in POLICES_CANDIDATS:
        try:
            font = ImageFont.truetype(chemin, TAILLE_PT)
            break
        except (IOError, OSError):
            continue
    if font is None:
        font = ImageFont.load_default()

    # Mesurer la séquence entière (pas caractère par caractère) pour éviter les gaps artificiels
    img_test = Image.new("RGB", (1, 1))
    draw_test = ImageDraw.Draw(img_test)
    bbox_seq = draw_test.textbbox((0, 0), sequence, font=font)
    seq_w = bbox_seq[2] - bbox_seq[0]
    seq_h = bbox_seq[3] - bbox_seq[1]

    largeur = PADDING * 2 + seq_w
    hauteur = PADDING * 2 + seq_h + 30  # +30 pour le label

    img = Image.new("RGB", (largeur, hauteur), FOND)
    draw = ImageDraw.Draw(img)

    draw.rectangle([(0, 0), (largeur - 1, hauteur - 1)], outline=BORDURE, width=3)

    try:
        font_label = ImageFont.truetype(POLICES_CANDIDATS[0], 13)
    except Exception:
        font_label = ImageFont.load_default()
    draw.text((PADDING, PADDING // 2), "Saisissez exactement :", font=font_label, fill=(100, 100, 100))

    # Séquence rendue d'un bloc — pas d'espacement artificiel entre les caractères
    y_texte = PADDING + 16
    draw.text((PADDING, y_texte), sequence, font=font, fill=TEXTE)

    # Soulignement continu sous toute la séquence pour montrer qu'elle est d'un seul tenant
    y_souligne = y_texte + seq_h + 3
    draw.line([(PADDING, y_souligne), (PADDING + seq_w, y_souligne)], fill=BORDURE, width=3)

    buf = io.BytesIO()
    img.save(buf, format="PNG", dpi=(150, 150))
    buf.seek(0)
    return buf


def _ajouter_paragraphe_instruction(doc, texte):
    """Ajoute une ligne de consigne en gris italique."""
    p = doc.add_paragraph()
    run = p.add_run(f"[ Consigne : {texte} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(9)
    return p


# ---------------------------------------------------------------------------
# Générateur principal
# ---------------------------------------------------------------------------

def generer_epreuve_word(
    contexte, session, annee, variante,
    competences_actives, points_par_competence=None,
    seed=None, competences=None, options_par_competence=None,
):
    """Génère (doc_etudiant, doc_corrige, config) pour une épreuve Word."""
    competences = competences or charger_competences_word()
    competences_par_id = {c["id"]: dict(c) for c in competences}
    for comp_id, pts in (points_par_competence or {}).items():
        if comp_id in competences_par_id:
            competences_par_id[comp_id]["points_defaut"] = pts

    actives = set(competences_actives) & set(competences_par_id)
    if not actives:
        raise ValueError("Aucune compétence active sélectionnée")

    ctx = CONTEXTES.get(contexte, CONTEXTES["sport"])
    doc_e = Document()
    doc_c = Document()

    criteres = []
    consignes_map = {}  # comp_id → texte consigne (assemblées dans l'ordre pédagogique à la fin)

    # Ordre pédagogique : structure physique → styles → formatage → éléments sémantiques → TDM en dernier
    ORDRE_PEDAGOGIQUE = [
        "mise_en_page",
        "structure_hierarchique",
        "styles_paragraphe",
        "mise_en_forme_caracteres",
        "rechercher_remplacer",
        "entete_pied_page",
        "saisie_texte",
        "stabilite_mise_en_page",
        "note_de_bas_de_page",
        "image",
        "table_des_matieres",
    ]

    # -----------------------------------------------------------------------
    # Mise en page (orientation + marges) — appliquée au corrigé uniquement
    # -----------------------------------------------------------------------
    orientation_attendue = "paysage"
    marge_cm = 2.0

    if "mise_en_page" in actives:
        for section in doc_c.sections:
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            section.left_margin = Cm(marge_cm)
            section.right_margin = Cm(marge_cm)
            section.top_margin = Cm(marge_cm)
            section.bottom_margin = Cm(marge_cm)

        consignes_map["mise_en_page"] = (
            f"Mise en page : passez le document en orientation paysage "
            f"et fixez les quatre marges à {marge_cm} cm."
        )
        criteres.append({
            "id": "mise_en_page", "type": "mise_en_page",
            "competence": competences_par_id["mise_en_page"]["label"],
            "description": "Orientation paysage + marges 2 cm",
            "orientation": orientation_attendue, "marge_cm": marge_cm,
            "points": _points(competences_par_id, "mise_en_page"),
        })

    # -----------------------------------------------------------------------
    # En-tête et pied de page
    # -----------------------------------------------------------------------
    code_annee = str(annee)
    mots_cles_entete = [session, code_annee]
    mots_cles_pied = ["confidentiel"]

    if "entete_pied_page" in actives:
        # Corrigé : remplit en-tête et pied de page
        entete_c = doc_c.sections[0].header
        entete_c.paragraphs[0].text = f"{session} {annee} — Document officiel"
        pied_c = doc_c.sections[0].footer
        pied_c.paragraphs[0].text = "Document confidentiel"

        consignes_map["entete_pied_page"] = (
            f"En-tête : saisissez \"{session} {annee} — Document officiel\". "
            f"Pied de page : saisissez \"Document confidentiel\"."
        )
        criteres.append({
            "id": "entete_pied_page", "type": "entete_pied_page",
            "competence": competences_par_id["entete_pied_page"]["label"],
            "description": "En-tête et pied de page renseignés",
            "mots_cles_entete": mots_cles_entete,
            "mots_cles_pied": mots_cles_pied,
            "points": _points(competences_par_id, "entete_pied_page"),
        })

    # -----------------------------------------------------------------------
    # Titre principal
    # -----------------------------------------------------------------------
    ANCRE_TITRE = ctx["titre"]
    for doc, style in ((doc_e, "Normal"), (doc_c, "Heading 1")):
        p = doc.add_paragraph(ANCRE_TITRE, style=style)

    # -----------------------------------------------------------------------
    # Table des matières (placeholder dans étudiant, champ dans corrigé)
    # -----------------------------------------------------------------------
    if "table_des_matieres" in actives:
        _ajouter_paragraphe_instruction(doc_e, "Insérez ici une table des matières automatique.")
        toc_c = doc_c.add_paragraph()
        _inserer_champ_toc(toc_c)

        consignes_map["table_des_matieres"] = (
            "Table des matières : insérez une table des matières automatique "
            "après avoir appliqué les styles de titres (cette étape doit être réalisée en dernier)."
        )
        criteres.append({
            "id": "table_des_matieres", "type": "table_des_matieres",
            "competence": competences_par_id["table_des_matieres"]["label"],
            "description": "Table des matières générée automatiquement",
            "points": _points(competences_par_id, "table_des_matieres"),
        })

    # -----------------------------------------------------------------------
    # Section 1 — Introduction
    # -----------------------------------------------------------------------
    ANCRE_INTRO = "Introduction"
    ANCRE_INTRO_SS = "Présentation du secteur"

    for doc, h1, h2 in (
        (doc_e, "Normal", "Normal"),
        (doc_c, "Heading 1", "Heading 2"),
    ):
        doc.add_paragraph(ANCRE_INTRO, style=h1)
        doc.add_paragraph(ANCRE_INTRO_SS, style=h2)

    # Corps d'introduction (avec mot à formater)
    ANCRE_FORMAT = ctx["mot_formater"]
    intro_text = ctx["intro_body"]
    for doc in (doc_e, doc_c):
        doc.add_paragraph(intro_text)

    # Mise en forme des caractères dans le corrigé
    if "mise_en_forme_caracteres" in actives:
        # Reformate le paragraphe dans le corrigé avec le mot en gras + couleur
        para_c = doc_c.paragraphs[-1]
        para_c.clear()
        before, _, after = intro_text.partition(ANCRE_FORMAT)
        para_c.add_run(before)
        run_fmt = para_c.add_run(ANCRE_FORMAT)
        _set_run_format(run_fmt, gras=True, taille_pt=13, couleur_hex="E0007A")
        para_c.add_run(after)

        note_conflit = (
            " Attention : appliquez d'abord le style « Corps de texte » (si demandé), "
            "puis effectuez cette mise en forme."
            if "styles_paragraphe" in actives else ""
        )
        consignes_map["mise_en_forme_caracteres"] = (
            f"Mise en forme : dans le premier paragraphe de l'introduction, "
            f"mettez le mot « {ANCRE_FORMAT} » en gras, taille 13, couleur magenta "
            f"(code hex #E0007A — ou via Autres couleurs > Personnalisée : R:224 V:0 B:122)."
            f"{note_conflit}"
        )
        criteres.append({
            "id": "mise_en_forme_caracteres", "type": "mise_en_forme_caracteres",
            "competence": competences_par_id["mise_en_forme_caracteres"]["label"],
            "description": f"Mot « {ANCRE_FORMAT} » : gras, 13 pt, couleur #E0007A",
            "ancre_paragraphe": intro_text[:40],
            "mot_cible": ANCRE_FORMAT,
            "gras": True, "taille_pt": 13, "couleur_hex": "E0007A",
            "points": _points(competences_par_id, "mise_en_forme_caracteres"),
        })

    # Rechercher-remplacer
    TERME_ANCIEN = ctx["terme_remplacer"]
    TERME_NOUVEAU = ctx["terme_nouveau"]

    if "rechercher_remplacer" in actives:
        consignes_map["rechercher_remplacer"] = (
            f"Rechercher-remplacer : remplacez toutes les occurrences de "
            f"« {TERME_ANCIEN} » par « {TERME_NOUVEAU} » dans tout le document."
        )
        criteres.append({
            "id": "rechercher_remplacer", "type": "rechercher_remplacer",
            "competence": competences_par_id["rechercher_remplacer"]["label"],
            "description": f"Remplacer « {TERME_ANCIEN} » par « {TERME_NOUVEAU} »",
            "terme_ancien": TERME_ANCIEN,
            "terme_nouveau": TERME_NOUVEAU,
            "points": _points(competences_par_id, "rechercher_remplacer"),
        })

    # Saisie de texte — séquence de caractères spéciaux affichée en image (non copiable)
    ANCRE_SAISIE = "Zone de saisie"

    # Séquences par variante — différentes pour empêcher la copie entre étudiants
    SEQUENCES = [
        "{ABC}[€2025]",
        "[XYZ]@{€100}",
        "€{DEF}[50%]",
        "{GHI}[€75]@",
        "[JKL]{€30}%",
        "€[MNO]{25%}",
    ]
    idx_variante = ord(variante.upper()) - ord("A") if variante else 0
    SEQ_ATTENDUE = SEQUENCES[idx_variante % len(SEQUENCES)]
    CHARS_SPECIAUX = sorted(set(c for c in SEQ_ATTENDUE if not c.isalnum()))

    if "saisie_texte" in actives:
        _ajouter_paragraphe_instruction(
            doc_e,
            "Lisez la séquence affichée dans l'image ci-dessous et saisissez-la "
            "exactement dans la zone de saisie, caractère par caractère au clavier."
        )

        # --- Génération de l'image (Pillow) ---
        img_buf = _generer_image_sequence(SEQ_ATTENDUE)
        doc_e.add_picture(img_buf, width=Cm(10))
        doc_e.add_paragraph("")

        # Ancre (invisible pour l'étudiant, sert au correcteur)
        p_ancre = doc_e.add_paragraph(ANCRE_SAISIE)
        p_ancre.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # blanc = invisible
        p_ancre.runs[0].font.size = Pt(1)

        # Tableau de saisie : ligne 0 = label, ligne 1 = zone de frappe
        _ajouter_tableau_saisie(doc_e)

        # Corrigé : ancre + tableau avec séquence correcte
        doc_c.add_paragraph(ANCRE_SAISIE)
        _ajouter_tableau_saisie(doc_c, texte_prefill=SEQ_ATTENDUE)

        consignes_map["saisie_texte"] = (
            "Saisie de texte : observez l'image ci-dessous et recopiez "
            "la séquence exactement dans la zone prévue, "
            "caractère par caractère au clavier."
        )
        criteres.append({
            "id": "saisie_texte", "type": "saisie_texte",
            "competence": competences_par_id["saisie_texte"]["label"],
            "description": f"Séquence {SEQ_ATTENDUE!r} saisie depuis une image",
            "ancre_zone": ANCRE_SAISIE,
            "sequence_attendue": SEQ_ATTENDUE,
            "chars_speciaux": CHARS_SPECIAUX,
            "points": _points(competences_par_id, "saisie_texte"),
        })

    # -----------------------------------------------------------------------
    # Section 2 — Données
    # -----------------------------------------------------------------------
    ANCRE_DONNEES = "Données et résultats"
    ANCRE_DONNEES_SS = "Analyse statistique"

    for doc, h1, h2 in (
        (doc_e, "Normal", "Normal"),
        (doc_c, "Heading 1", "Heading 2"),
    ):
        doc.add_paragraph(ANCRE_DONNEES, style=h1)
        doc.add_paragraph(ANCRE_DONNEES_SS, style=h2)
        doc.add_paragraph(ctx["donnees_body"])

    # Image
    ANCRE_IMAGE = "Insérez ici une image illustrant le rapport"
    if "image" in actives:
        doc_e.add_paragraph(ANCRE_IMAGE)
        _ajouter_paragraphe_instruction(
            doc_e, "Insérez une image (largeur entre 5 et 12 cm)."
        )
        doc_c.add_paragraph("[Image insérée par l'enseignant dans le corrigé]")

        consignes_map["image"] = (
            "Image : insérez une image illustrant le rapport. "
            "Sa largeur doit être comprise entre 5 et 12 cm."
        )
        criteres.append({
            "id": "image", "type": "image",
            "competence": competences_par_id["image"]["label"],
            "description": "Image insérée, largeur 5–12 cm",
            "largeur_min_cm": 5.0, "largeur_max_cm": 12.0,
            "points": _points(competences_par_id, "image"),
        })

    # Stabilité de la mise en page (saut de page)
    ANCRE_SAUT = "Suite du document — nouvelle page"
    if "stabilite_mise_en_page" in actives:
        doc_e.add_paragraph(ANCRE_SAUT)
        _ajouter_paragraphe_instruction(
            doc_e, f"Insérez un saut de page avant le paragraphe « {ANCRE_SAUT} »."
        )
        para_c_saut = doc_c.add_paragraph(ANCRE_SAUT)
        _inserer_saut_de_page(para_c_saut)

        consignes_map["stabilite_mise_en_page"] = (
            f"Stabilité de la mise en page : insérez un saut de page "
            f"avant le paragraphe « {ANCRE_SAUT} »."
        )
        criteres.append({
            "id": "stabilite_mise_en_page", "type": "stabilite_mise_en_page",
            "competence": competences_par_id["stabilite_mise_en_page"]["label"],
            "description": "Saut de page avant la section 2",
            "ancre_saut": ANCRE_SAUT,
            "points": _points(competences_par_id, "stabilite_mise_en_page"),
        })

    # Note de bas de page
    TERME_NOTE = ctx["terme_note"]
    ANCRE_NOTE = f"Le terme {TERME_NOTE} revêt une importance particulière dans ce domaine."

    if "note_de_bas_de_page" in actives:
        for doc in (doc_e, doc_c):
            doc.add_paragraph(ANCRE_NOTE)
        _ajouter_paragraphe_instruction(
            doc_e,
            f"Ajoutez une note de bas de page sur le terme « {TERME_NOTE} » "
            f"dans le paragraphe ci-dessus."
        )

        consignes_map["note_de_bas_de_page"] = (
            f"Note de bas de page : dans le paragraphe mentionnant « {TERME_NOTE} », "
            f"ajoutez une note de bas de page sur ce terme."
        )
        criteres.append({
            "id": "note_de_bas_de_page", "type": "note_de_bas_de_page",
            "competence": competences_par_id["note_de_bas_de_page"]["label"],
            "description": f"Note de bas de page sur « {TERME_NOTE} »",
            "nb_notes_min": 1,
            "points": _points(competences_par_id, "note_de_bas_de_page"),
        })

    # -----------------------------------------------------------------------
    # Section 3 — Conclusion
    # -----------------------------------------------------------------------
    ANCRE_CONCLUSION = "Conclusion"
    for doc, h1 in ((doc_e, "Normal"), (doc_c, "Heading 1")):
        doc.add_paragraph(ANCRE_CONCLUSION, style=h1)
    for doc in (doc_e, doc_c):
        doc.add_paragraph(
            "Ce rapport présente une synthèse des éléments analysés. "
            "Les résultats obtenus permettent de dégager des perspectives d'amélioration."
        )

    # -----------------------------------------------------------------------
    # Structure hiérarchique (critère portant sur les titres)
    # -----------------------------------------------------------------------
    para_titres = [
        {"ancre": ANCRE_TITRE, "style": "heading1"},
        {"ancre": ANCRE_INTRO, "style": "heading1"},
        {"ancre": ANCRE_INTRO_SS, "style": "heading2"},
        {"ancre": ANCRE_DONNEES, "style": "heading1"},
        {"ancre": ANCRE_DONNEES_SS, "style": "heading2"},
        {"ancre": ANCRE_CONCLUSION, "style": "heading1"},
    ]
    if "structure_hierarchique" in actives:
        consignes_map["structure_hierarchique"] = (
            f"Structure hiérarchique : appliquez le style Titre 1 aux titres principaux "
            f"(\"{ANCRE_TITRE}\", Introduction, Données et résultats, Conclusion) "
            "et Titre 2 aux sous-titres (Présentation du secteur, Analyse statistique)."
        )
        criteres.append({
            "id": "structure_hierarchique", "type": "structure_hierarchique",
            "competence": competences_par_id["structure_hierarchique"]["label"],
            "description": "Styles Titre 1 / Titre 2 appliqués aux titres",
            "paragraphes": para_titres,
            "points": _points(competences_par_id, "structure_hierarchique"),
        })

    # Styles de paragraphe (corps de texte)
    if "styles_paragraphe" in actives:
        opts_sp = (options_par_competence or {}).get("styles_paragraphe", {})
        taille_imposee = opts_sp.get("taille_imposee", True)
        justifie = opts_sp.get("justifie", True)

        # Taille varie selon la variante (cycle 12/13/14 pt) pour empêcher la copie
        TAILLES_CYCLE = [13, 14, 12]
        idx_variante = ord(variante.upper()) - ord("A") if variante else 0
        TAILLE_CORPS = TAILLES_CYCLE[idx_variante % len(TAILLES_CYCLE)] if taille_imposee else None

        corps_ancres = [intro_text[:30], ctx["donnees_body"][:30]]
        for ancre in corps_ancres:
            for p in doc_c.paragraphs:
                if p.text.startswith(ancre.strip()):
                    try:
                        p.style = doc_c.styles["Body Text"]
                    except KeyError:
                        pass
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            style_corps = doc_c.styles["Body Text"]
            if taille_imposee:
                style_corps.font.size = Pt(TAILLE_CORPS)
            if justifie:
                style_corps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except (KeyError, Exception):
            pass

        # Construction de la consigne selon les options actives
        modifications = []
        if taille_imposee:
            modifications.append(f"taille de police {TAILLE_CORPS} pt")
        if justifie:
            modifications.append("texte justifié")

        if modifications:
            consigne_sp = (
                f"Styles de paragraphe : appliquez le style « Corps de texte » "
                f"aux paragraphes de contenu (paragraphes qui ne sont pas des titres), "
                f"puis modifiez ce style : {', '.join(modifications)}."
            )
        else:
            consigne_sp = (
                "Styles de paragraphe : appliquez le style « Corps de texte » "
                "aux paragraphes de contenu (paragraphes qui ne sont pas des titres)."
            )
        consignes_map["styles_paragraphe"] = consigne_sp

        critere_sp = {
            "id": "styles_paragraphe", "type": "styles_paragraphe",
            "competence": competences_par_id["styles_paragraphe"]["label"],
            "description": f"Style Corps de texte appliqué" + (f", modifié ({TAILLE_CORPS} pt, justifié)" if modifications else ""),
            "style_attendu": "body_text",
            "ancres": corps_ancres,
            "points": _points(competences_par_id, "styles_paragraphe"),
        }
        if taille_imposee:
            critere_sp["taille_pt"] = TAILLE_CORPS
        if justifie:
            critere_sp["justifie"] = True
        criteres.append(critere_sp)

    # -----------------------------------------------------------------------
    # Assemblage des consignes dans l'ordre pédagogique
    # -----------------------------------------------------------------------
    consignes = [consignes_map[k] for k in ORDRE_PEDAGOGIQUE if k in consignes_map]

    # -----------------------------------------------------------------------
    # Config JSON
    # -----------------------------------------------------------------------
    bareme = sum(c["points"] for c in criteres)
    exercice = {
        "id": "ex_word", "titre": "Document Word",
        "feuille": None,
        "criteres": criteres, "consignes": consignes,
    }
    config = {
        "session": session, "annee": annee, "module": "word",
        "variante": variante, "contexte": contexte,
        "exercices": [exercice],
        "bareme_total": bareme,
    }

    # -----------------------------------------------------------------------
    # Page de garde (étudiant) — insérée au début une fois toutes les consignes connues
    # -----------------------------------------------------------------------
    _inserer_page_garde_au_debut(doc_e, session, annee, variante, contexte, consignes)

    return doc_e, doc_c, config


def _inserer_page_garde_au_debut(doc, session, annee, variante, contexte, consignes):
    """Crée une page de garde avec toutes les consignes et l'insère en tête du document."""
    from docx import Document as _Doc
    tmp = _Doc()

    p_titre = tmp.add_paragraph()
    run = p_titre.add_run(f"Épreuve Word — {session} {annee} — Variante {variante}")
    run.bold = True
    run.font.size = Pt(14)

    tmp.add_paragraph(f"Contexte : {contexte.replace('_', ' ').capitalize()}")
    tmp.add_paragraph("Nom : _______________________   Prénom : _______________________")
    tmp.add_paragraph("")
    p_cons = tmp.add_paragraph()
    p_cons.add_run("Consignes de l'épreuve :").bold = True

    for i, consigne in enumerate(consignes, 1):
        tmp.add_paragraph(f"{i}. {consigne}")

    tmp.add_page_break()

    # Insère les paragraphes du document temporaire au début du vrai document
    import copy
    body = doc.element.body
    children = list(body)
    # Premier élément enfant du body (premier paragraphe ou sectPr)
    ref = children[0] if children else None
    insert_idx = 0
    for p in list(tmp.element.body):
        tag = p.tag.split("}")[-1] if "}" in p.tag else p.tag
        if tag in ("p", "tbl"):
            clone = copy.deepcopy(p)
            body.insert(insert_idx, clone)
            insert_idx += 1
